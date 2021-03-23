#!/usr/bin/env python3.7
from pdfminer.layout import LAParams, LTTextBox
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt
import tkinter as tk

sku = []
fnsku = []
pieces = []
units = []
cases = []
total = []

# Possible x-ranges: units, cases, total
xrange = [
    [[400, 450], [480, 500], [510, 530]],  # ZYHT
    [[450, 470], [480, 500], [508, 520]],  # NPHR
    [[450, 490], [520, 530], [530, 600]]  # Q3H0
]


def read_pdf(file, range_set, output_file):
    sku = []
    fnsku = []
    pieces = []
    units = []
    cases = []
    total = []
    print(file)
    fp = open(file, 'rb')
    rsrcmgr = PDFResourceManager()
    laparams = LAParams()
    device = PDFPageAggregator(rsrcmgr, laparams=laparams)
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    pages = PDFPage.get_pages(fp)

    for page in pages:
        print('Processing next page...')
        interpreter.process_page(page)
        layout = device.get_result()
        for lobj in layout:
            if isinstance(lobj, LTTextBox):
                x, y, text = lobj.bbox[0], lobj.bbox[3], lobj.get_text()  # 595 W X 842 H pixels
                # print('At %r is text: %s' % ((x, y), text))
                output_file.write('At %r is text: %s' % ((x, y), text))

                # sku
                import re
                if int(x) in range(35, 50):
                    if len(text) >= 6:
                        sku.append(text)
                        seg = text.split('-')
                        pieces.append(re.sub("[^0-9]", "", seg[1]))
                        # print('At %r is text: %s' % ((x, y), text))

                # fnsku
                if int(x) in range(85, 120):
                    if text[-11:][0:2] == 'X0' or text[-11:][0:2] == 'B0':
                        fnsku.append(text[-11:])
                        # print('At %r is text: %s' % ((x, y), text))

                # If SKU does not contain pieces-info, use below
                #
                # if 'pack' in text:
                #     pos = text.index('pack')
                #     pieces.append(text[pos-3:pos])
                # elif 'pcs' in text:
                #     pos = text.index('pcs')
                #     pieces.append(text[pos-3:pos])
                # elif 'pack of' in text:
                #     pos = text.index('pack of')
                #     pieces.append(text[pos + 8:pos + 9])
                # else:
                #     pieces.append(1)

                # units
                if int(x) in range(xrange[range_set][0][0], xrange[range_set][0][1]):
                    units.append(text)
                    # print('At %r is text: %s' % ((x, y), text))

                # cases
                if int(x) in range(xrange[range_set][1][0], xrange[range_set][1][1]):
                    cases.append(text)
                    # print('At %r is text: %s' % ((x, y), text))

                # total
                if int(x) in range(xrange[range_set][2][0], xrange[range_set][2][1]):
                    total.append(text)
                    # print('At %r is text: %s' % ((x, y), text))
    fp.close()
    return sku, fnsku, pieces, units, cases, total


def make_doc(order, sku, fnsku, pieces, units, cases, total):
    document = Document()
    # print(sku)
    # print(fnsku)
    # print(units)
    # print(cases)
    # print(total)

    for i in range(len(sku)):
        run = document.add_paragraph().add_run(order)
        run.font.size = Pt(24)
        temp_sku = sku[i].replace('\n', '') + '\n'
        temp_fnsku = fnsku[i].replace('\n', '') + '\n'
        temp_units = pieces[i].replace('\n', '') + ' PCS/UNIT\n'
        temp_cases = str(int(int(units[i].replace('\n', '')) / int(cases[i].replace('\n', '')))) + ' UNITS/CASE\n'
        temp_total = 'TOTAL: ' + cases[i].replace('\n', '') + ' CASES, ' + units[i].replace('\n', '') + ' UNITS'
        text = temp_sku + temp_fnsku + temp_units + temp_cases + temp_total
        run = document.add_paragraph().add_run(text)
        run.font.size = Pt(24)
        run.add_break(WD_BREAK.PAGE)

    document.save('../../shipment.docx')  # save in main production dir
    print("Finished Processing. Ok to close.")


def main():
    root = tk.Tk()
    frame = tk.Frame(master=root, width=300, height=300)
    frame.grid(row=4, column=2)
    frame.pack()

    info = "Note: place pdf file in same folder as program"
    info_label = tk.Label(master=frame, text=info)
    info_label.grid(row=1, column=2)

    v1 = tk.StringVar()
    label1 = tk.Label(master=frame, text='Shipment Number: ')
    entry1 = tk.Entry(master=frame, width=50, textvariable=v1)
    label1.grid(row=2, column=1)
    entry1.grid(row=2, column=2)

    v2 = tk.StringVar()
    label2 = tk.Label(master=frame, text='Filename of shipment PDF with extension: ')
    entry2 = tk.Entry(master=frame, width=50, textvariable=v2)
    label2.grid(row=3, column=1)
    entry2.grid(row=3, column=2)

    def update():
        order = v1.get()
        file = "../../" + v2.get()  # exe --> dist --> src --> production
        output_file = open("debug_print.txt", "w+")
        checking = True
        range_set = 0  # for xrange
        sku, fnsku, pieces, units, cases, total = read_pdf(file, range_set, output_file)

        # PDF format varies, check other sets of x-positions of text boxes
        running = True
        while running:
            if not units or not cases or not total or 'New\n' in units:
                range_set += 1
                sku, fnsku, pieces, units, cases, total = read_pdf(file, range_set, output_file)
                if 'New\n' not in units and units and cases and total:
                    running = False
            else:
                running = False

        make_doc(order, sku, fnsku, pieces, units, cases, total)
        output_file.close()

    button = tk.Button(master=frame, text='Make Doc', command=update)
    button.grid(row=4, column=2)

    root.mainloop()


main()

